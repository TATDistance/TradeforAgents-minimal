#!/usr/bin/env python3
"""
DeepSeek 单模型多模块股票分析工具（CLI）

目标：
- 仅输入股票代码（可选日期）
- 仅用 DeepSeek API
- 产出接近原项目的多报告目录结构

用法:
  python scripts/minimal_deepseek_report.py 600028
  python scripts/minimal_deepseek_report.py 510300 --date 2026-03-26
  python scripts/minimal_deepseek_report.py AAPL --model deepseek-chat

输出目录:
  results/<输入代码>/<日期>/
    - analysis_metadata.json
    - decision.json
    - message_tool.log
    - reports/
      - market_report.md
      - fundamentals_report.md
      - news_report.md
      - research_team_decision.md
      - investment_plan.md
      - trader_investment_plan.md
      - risk_management_decision.md
      - final_trade_decision.md
      - final_report.md
      - market_snapshot.json
      - news_snapshot.json
"""

from __future__ import annotations

import argparse
import contextlib
import html
import json
import logging
import os
import re
import time
import urllib.parse
import urllib.request
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any, Dict, List, Tuple

import pandas as pd
import yfinance as yf
from openai import OpenAI


# yfinance 会输出大量 warning（包含非致命网络重试），统一降级避免误判为失败
logging.getLogger("yfinance").setLevel(logging.ERROR)

EASTMONEY_HEADERS = {
    "User-Agent": "Mozilla/5.0",
    "Referer": "https://quote.eastmoney.com/",
}

EASTMONEY_UT = "7eea3edcaed734bea9cbfc24409ed989"


def normalize_symbol(user_symbol: str) -> Tuple[str, str]:
    """返回 (原始输入, yfinance可用代码)"""
    raw = user_symbol.strip().upper()
    if not raw:
        raise ValueError("股票代码不能为空")

    if re.fullmatch(r"[A-Z0-9.\-]{1,20}", raw) and "." in raw:
        return raw, raw

    if re.fullmatch(r"\d{6}", raw):
        # 上交所常见前缀: 5(ETF/基金), 6(主板/科创), 9(B股)
        if raw.startswith(("5", "6", "9")):
            return raw, f"{raw}.SS"
        return raw, f"{raw}.SZ"

    if re.fullmatch(r"\d{4,5}", raw):
        hk_code = str(int(raw)).zfill(4)
        return raw, f"{hk_code}.HK"

    if re.fullmatch(r"[A-Z][A-Z0-9\-]{0,9}", raw):
        return raw, raw

    raise ValueError(f"不支持的股票代码格式: {user_symbol}")


def infer_market_type(yf_symbol: str) -> str:
    if yf_symbol.endswith(".SS") or yf_symbol.endswith(".SZ"):
        return "A股"
    if yf_symbol.endswith(".HK"):
        return "港股"
    return "美股"


def candidate_yf_symbols(user_symbol: str, preferred_yf_symbol: str) -> List[str]:
    cands: List[str] = [preferred_yf_symbol]
    if re.fullmatch(r"\d{6}", user_symbol):
        if user_symbol.startswith(("5", "6", "9")):
            cands.extend([f"{user_symbol}.SS"])
        elif user_symbol.startswith(("0", "1", "2", "3")):
            cands.extend([f"{user_symbol}.SZ"])
        else:
            cands.extend([f"{user_symbol}.SS", f"{user_symbol}.SZ"])

    seen = set()
    uniq: List[str] = []
    for c in cands:
        if c not in seen:
            seen.add(c)
            uniq.append(c)
    return uniq


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def _safe_cache_key(*parts: str) -> str:
    return "__".join(re.sub(r"[^A-Za-z0-9._-]+", "_", p) for p in parts)


def read_json_cache(cache_file: Path, ttl_seconds: int) -> Dict[str, Any] | None:
    if not cache_file.exists():
        return None
    if ttl_seconds > 0:
        age = time.time() - cache_file.stat().st_mtime
        if age > ttl_seconds:
            return None
    try:
        return json.loads(cache_file.read_text(encoding="utf-8"))
    except Exception:
        return None


def write_json_cache(cache_file: Path, payload: Any) -> None:
    ensure_dir(cache_file.parent)
    cache_file.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def load_latest_market_cache_snapshot(cache_dir: Path, yf_symbol: str) -> Dict[str, Any] | None:
    pattern = f"market__{re.sub(r'[^A-Za-z0-9._-]+', '_', yf_symbol)}__*.json"
    candidates = sorted(cache_dir.glob(pattern), key=lambda item: item.stat().st_mtime, reverse=True)
    for path in candidates:
        try:
            return json.loads(path.read_text(encoding="utf-8"))
        except Exception:
            continue
    return None


def load_recent_result_market_snapshot(results_dir: Path, user_symbol: str, analysis_date: str) -> Dict[str, Any] | None:
    symbol_root = results_dir / user_symbol
    if not symbol_root.exists():
        return None

    def _sort_key(path: Path) -> str:
        return path.parent.parent.name

    candidates = sorted(
        symbol_root.glob("*/reports/market_snapshot.json"),
        key=_sort_key,
        reverse=True,
    )
    for path in candidates:
        try:
            snap_date = path.parent.parent.name
            if snap_date > analysis_date:
                continue
            return json.loads(path.read_text(encoding="utf-8"))
        except Exception:
            continue
    return None


def pct_change(series: pd.Series, periods: int) -> float | None:
    if len(series) <= periods:
        return None
    old = float(series.iloc[-periods - 1])
    new = float(series.iloc[-1])
    if old == 0:
        return None
    return (new - old) / old


def is_a_share_symbol(yf_symbol: str) -> bool:
    return bool(re.fullmatch(r"\d{6}\.(SS|SZ)", yf_symbol))


def eastmoney_secid(yf_symbol: str) -> str:
    code, market = yf_symbol.split(".", 1)
    if market == "SS":
        return f"1.{code}"
    return f"0.{code}"


def eastmoney_request_json(url: str, params: Dict[str, Any], timeout: int = 15) -> Dict[str, Any]:
    query = urllib.parse.urlencode(params)
    request = urllib.request.Request(f"{url}?{query}", headers=EASTMONEY_HEADERS)
    opener = urllib.request.build_opener(urllib.request.ProxyHandler({}))
    with opener.open(request, timeout=timeout) as response:
        return json.loads(response.read().decode("utf-8"))


def fetch_market_data_eastmoney(yf_symbol: str, as_of_date: str) -> Dict[str, Any]:
    end_dt = datetime.strptime(as_of_date, "%Y-%m-%d")
    start_dt = end_dt - timedelta(days=220)
    code = yf_symbol.split(".", 1)[0]

    payload = eastmoney_request_json(
        "https://push2his.eastmoney.com/api/qt/stock/kline/get",
        params={
            "secid": eastmoney_secid(yf_symbol),
            "ut": EASTMONEY_UT,
            "fields1": "f1,f2,f3,f4,f5,f6",
            "fields2": "f51,f52,f53,f54,f55,f56,f57,f58,f59,f60,f61",
            "klt": "101",
            "fqt": "1",
            "beg": start_dt.strftime("%Y%m%d"),
            "end": end_dt.strftime("%Y%m%d"),
        },
        timeout=15,
    )
    data = payload.get("data") or {}
    klines = data.get("klines") or []
    if not klines:
        raise RuntimeError(f"无法从东财获取 A 股行情: {yf_symbol}")

    rows = []
    for line in klines:
        parts = str(line).split(",")
        if len(parts) < 7:
            continue
        rows.append(
            {
                "Date": parts[0],
                "Open": float(parts[1]),
                "Close": float(parts[2]),
                "High": float(parts[3]),
                "Low": float(parts[4]),
                "Volume": float(parts[5]),
                "Amount": float(parts[6]),
            }
        )
    if not rows:
        raise RuntimeError(f"东财返回的行情结构异常: {yf_symbol}")

    hist = pd.DataFrame(rows)
    close = hist["Close"].dropna()
    volume = hist["Volume"].dropna()
    latest = hist.iloc[-1]

    ma20 = float(close.tail(20).mean()) if len(close) >= 20 else None
    ma60 = float(close.tail(60).mean()) if len(close) >= 60 else None
    vol_avg20 = float(volume.tail(20).mean()) if len(volume) >= 20 else None

    p1w = pct_change(close, 5)
    p1m = pct_change(close, 21)
    p3m = pct_change(close, 63)
    high_52w = float(close.tail(252).max()) if len(close) >= 1 else None
    low_52w = float(close.tail(252).min()) if len(close) >= 1 else None

    # 基本面字段先尽量复用 yfinance 的 info，失败时允许为空
    info: Dict[str, Any] = {}
    try:
        info = yf.Ticker(yf_symbol).info or {}
    except Exception:
        info = {}

    return {
        "symbol": yf_symbol,
        "as_of_date": as_of_date,
        "data_source": "eastmoney_a_share",
        "latest": {
            "date": str(latest["Date"]),
            "open": float(latest["Open"]),
            "high": float(latest["High"]),
            "low": float(latest["Low"]),
            "close": float(latest["Close"]),
            "volume": float(latest["Volume"]),
            "amount": float(latest["Amount"]),
        },
        "technical": {
            "ma20": ma20,
            "ma60": ma60,
            "volume_avg20": vol_avg20,
            "change_1w": p1w,
            "change_1m": p1m,
            "change_3m": p3m,
            "high_52w": high_52w,
            "low_52w": low_52w,
        },
        "fundamentals": {
            "longName": data.get("name") or info.get("longName") or code,
            "sector": info.get("sector"),
            "industry": info.get("industry"),
            "marketCap": info.get("marketCap"),
            "trailingPE": info.get("trailingPE"),
            "forwardPE": info.get("forwardPE"),
            "priceToBook": info.get("priceToBook"),
            "dividendYield": info.get("dividendYield"),
            "beta": info.get("beta"),
            "currency": info.get("currency") or "CNY",
        },
    }


def fetch_market_data(yf_symbol: str, as_of_date: str) -> Dict[str, Any]:
    if is_a_share_symbol(yf_symbol):
        try:
            return fetch_market_data_eastmoney(yf_symbol, as_of_date)
        except Exception as exc:
            raise RuntimeError("无法从东财获取 A 股行情: {0}".format(exc)) from exc

    end_dt = datetime.strptime(as_of_date, "%Y-%m-%d")
    start_dt = end_dt - timedelta(days=220)

    ticker = yf.Ticker(yf_symbol)
    hist = ticker.history(
        start=start_dt.strftime("%Y-%m-%d"),
        end=(end_dt + timedelta(days=1)).strftime("%Y-%m-%d"),
    )

    if hist is None or hist.empty:
        raise RuntimeError(
            "无法从 yfinance 获取行情: {0}（返回空数据，常见原因是 Yahoo 对当前 A 股代码无覆盖、接口抖动或网络不稳定）".format(
                yf_symbol
            )
        )

    close = hist["Close"].dropna()
    volume = hist["Volume"].dropna()

    latest = hist.iloc[-1]
    ma20 = float(close.tail(20).mean()) if len(close) >= 20 else None
    ma60 = float(close.tail(60).mean()) if len(close) >= 60 else None
    vol_avg20 = float(volume.tail(20).mean()) if len(volume) >= 20 else None

    p1w = pct_change(close, 5)
    p1m = pct_change(close, 21)
    p3m = pct_change(close, 63)

    high_52w = float(close.tail(252).max()) if len(close) >= 1 else None
    low_52w = float(close.tail(252).min()) if len(close) >= 1 else None

    info: Dict[str, Any] = {}
    try:
        info = ticker.info or {}
    except Exception:
        info = {}

    snapshot = {
        "symbol": yf_symbol,
        "as_of_date": as_of_date,
        "data_source": "yfinance",
        "latest": {
            "date": str(hist.index[-1].date()),
            "open": float(latest.get("Open", 0.0)),
            "high": float(latest.get("High", 0.0)),
            "low": float(latest.get("Low", 0.0)),
            "close": float(latest.get("Close", 0.0)),
            "volume": float(latest.get("Volume", 0.0)),
        },
        "technical": {
            "ma20": ma20,
            "ma60": ma60,
            "volume_avg20": vol_avg20,
            "change_1w": p1w,
            "change_1m": p1m,
            "change_3m": p3m,
            "high_52w": high_52w,
            "low_52w": low_52w,
        },
        "fundamentals": {
            "longName": info.get("longName"),
            "sector": info.get("sector"),
            "industry": info.get("industry"),
            "marketCap": info.get("marketCap"),
            "trailingPE": info.get("trailingPE"),
            "forwardPE": info.get("forwardPE"),
            "priceToBook": info.get("priceToBook"),
            "dividendYield": info.get("dividendYield"),
            "beta": info.get("beta"),
            "currency": info.get("currency"),
        },
    }
    return snapshot


def fetch_news(yf_symbol: str, max_news: int = 8) -> List[Dict[str, Any]]:
    ticker = yf.Ticker(yf_symbol)
    try:
        items = ticker.news or []
    except Exception:
        items = []

    parsed: List[Dict[str, Any]] = []
    for item in items[:max_news]:
        content = item.get("content", {}) if isinstance(item, dict) else {}
        parsed.append(
            {
                "title": content.get("title") or item.get("title"),
                "summary": content.get("summary") or item.get("summary"),
                "source": (
                    (content.get("provider") or {}).get("displayName")
                    if isinstance(content.get("provider"), dict)
                    else item.get("publisher")
                ),
                "url": (
                    content.get("canonicalUrl", {}).get("url")
                    if isinstance(content.get("canonicalUrl"), dict)
                    else item.get("link")
                ),
                "published_at": content.get("pubDate") or item.get("providerPublishTime"),
            }
        )
    return parsed


def fetch_market_data_cached(
    yf_symbol: str,
    as_of_date: str,
    cache_dir: Path,
    results_dir: Path,
    user_symbol: str,
    module_logs: List[str],
    cache_ttl_seconds: int,
    disable_cache: bool,
) -> Dict[str, Any]:
    cache_key = _safe_cache_key("market", yf_symbol, as_of_date)
    cache_file = cache_dir / f"{cache_key}.json"
    if not disable_cache:
        cached = read_json_cache(cache_file, ttl_seconds=cache_ttl_seconds)
        if cached is not None:
            module_logs.append(f"[{datetime.now().isoformat()}] market_cache hit file={cache_file.name}")
            return cached

    try:
        snapshot = fetch_market_data(yf_symbol, as_of_date)
    except Exception as exc:
        if cache_file.exists():
            module_logs.append(
                f"[{datetime.now().isoformat()}] market_cache stale_fallback file={cache_file.name} err={exc}"
            )
            return json.loads(cache_file.read_text(encoding="utf-8"))
        latest_cached = load_latest_market_cache_snapshot(cache_dir, yf_symbol)
        if latest_cached is not None:
            module_logs.append(
                f"[{datetime.now().isoformat()}] market_cache recent_symbol_fallback symbol={yf_symbol} err={exc}"
            )
            return latest_cached
        if is_a_share_symbol(yf_symbol):
            recent_snapshot = load_recent_result_market_snapshot(results_dir, user_symbol, as_of_date)
            if recent_snapshot is not None:
                module_logs.append(
                    f"[{datetime.now().isoformat()}] market_result_fallback symbol={user_symbol} err={exc}"
                )
                return recent_snapshot
        raise
    if not disable_cache:
        write_json_cache(cache_file, snapshot)
        module_logs.append(f"[{datetime.now().isoformat()}] market_cache write file={cache_file.name}")
    return snapshot


def fetch_news_cached(
    yf_symbol: str,
    max_news: int,
    cache_dir: Path,
    module_logs: List[str],
    cache_ttl_seconds: int,
    disable_cache: bool,
) -> List[Dict[str, Any]]:
    cache_key = _safe_cache_key("news", yf_symbol, f"top{max_news}")
    cache_file = cache_dir / f"{cache_key}.json"
    if not disable_cache:
        cached = read_json_cache(cache_file, ttl_seconds=cache_ttl_seconds)
        if isinstance(cached, list):
            module_logs.append(f"[{datetime.now().isoformat()}] news_cache hit file={cache_file.name}")
            return cached

    news = fetch_news(yf_symbol, max_news=max_news)
    if not disable_cache:
        write_json_cache(cache_file, news)
        module_logs.append(f"[{datetime.now().isoformat()}] news_cache write file={cache_file.name}")
    return news


def build_data_context(
    user_symbol: str,
    yf_symbol: str,
    analysis_date: str,
    market_type: str,
    market_snapshot: Dict[str, Any],
    news: List[Dict[str, Any]],
) -> str:
    return (
        f"用户输入代码: {user_symbol}\n"
        f"规范化代码: {yf_symbol}\n"
        f"市场类型: {market_type}\n"
        f"分析日期: {analysis_date}\n\n"
        "行情与基本面JSON:\n"
        f"{json.dumps(market_snapshot, ensure_ascii=False, indent=2)}\n\n"
        "新闻JSON:\n"
        f"{json.dumps(news, ensure_ascii=False, indent=2)}\n"
    )


REUSED_DIRECTION_REPORTS = [
    "fundamentals_report.md",
    "news_report.md",
    "research_team_decision.md",
    "investment_plan.md",
    "risk_management_decision.md",
]


def _load_json_file(path: Path) -> Dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def _parse_date(value: str) -> datetime.date | None:
    try:
        return datetime.strptime(value, "%Y-%m-%d").date()
    except Exception:
        return None


def _result_dir_is_reusable(result_dir: Path) -> bool:
    decision_path = result_dir / "decision.json"
    reports_dir = result_dir / "reports"
    if not decision_path.exists() or not reports_dir.exists():
        return False

    try:
        decision = _load_json_file(decision_path)
    except Exception:
        return False

    if decision.get("degraded_modules"):
        return False

    required = ["final_report.md", "final_trade_decision.md"] + REUSED_DIRECTION_REPORTS
    return all((reports_dir / name).exists() for name in required)


def find_recent_reusable_result(
    results_dir: Path,
    user_symbol: str,
    analysis_date: str,
    direction_cache_days: int,
) -> Dict[str, Any] | None:
    symbol_root = results_dir / user_symbol
    if not symbol_root.exists():
        return None

    target_date = _parse_date(analysis_date)
    if target_date is None:
        return None

    exact_match = None
    recent_match = None

    for child in symbol_root.iterdir():
        if not child.is_dir():
            continue
        child_date = _parse_date(child.name)
        if child_date is None:
            continue
        if not _result_dir_is_reusable(child):
            continue

        delta_days = (target_date - child_date).days
        if delta_days == 0:
            exact_match = {
                "type": "same_day",
                "result_dir": child,
                "analysis_date": child.name,
                "delta_days": delta_days,
            }
            break
        if 0 < delta_days <= direction_cache_days:
            if recent_match is None or child.name > str(recent_match["analysis_date"]):
                recent_match = {
                    "type": "direction_cache",
                    "result_dir": child,
                    "analysis_date": child.name,
                    "delta_days": delta_days,
                }

    return exact_match or recent_match


def load_reused_reports(result_dir: Path) -> Dict[str, str]:
    reports_dir = result_dir / "reports"
    reused = {}
    for name in REUSED_DIRECTION_REPORTS:
        path = reports_dir / name
        if path.exists():
            reused[name] = path.read_text(encoding="utf-8")
    return reused


def safe_json_extract(text: str) -> Dict[str, Any] | None:
    match = re.search(r"\{[\s\S]*\}", text)
    if not match:
        return None
    try:
        return json.loads(match.group(0))
    except json.JSONDecodeError:
        return None


def _get_field(obj: Any, key: str, default: Any = None) -> Any:
    if obj is None:
        return default
    if isinstance(obj, dict):
        return obj.get(key, default)
    return getattr(obj, key, default)


def _logical_module_name(module_name: str) -> str:
    return module_name[:-9] if module_name.endswith("_fallback") else module_name


def extract_usage_info(resp: Any) -> Dict[str, Any]:
    usage = _get_field(resp, "usage", {})
    prompt_tokens = _get_field(usage, "prompt_tokens", None)
    completion_tokens = _get_field(usage, "completion_tokens", None)
    total_tokens = _get_field(usage, "total_tokens", None)

    # 兼容部分 OpenAI-compatible 返回字段
    if prompt_tokens is None:
        prompt_tokens = _get_field(usage, "input_tokens", 0)
    if completion_tokens is None:
        completion_tokens = _get_field(usage, "output_tokens", 0)
    if total_tokens is None:
        total_tokens = int(prompt_tokens or 0) + int(completion_tokens or 0)

    choices = _get_field(resp, "choices", []) or []
    finish_reason = ""
    if choices:
        c0 = choices[0]
        finish_reason = str(_get_field(c0, "finish_reason", "") or "")

    return {
        "prompt_tokens": int(prompt_tokens or 0),
        "completion_tokens": int(completion_tokens or 0),
        "total_tokens": int(total_tokens or 0),
        "finish_reason": finish_reason,
    }


def summarize_module_metrics(module_metrics: List[Dict[str, Any]]) -> Dict[str, Any]:
    per_module: Dict[str, Dict[str, Any]] = {}
    total_prompt = 0
    total_completion = 0
    total_tokens = 0
    total_elapsed = 0.0
    success_calls = 0
    fail_calls = 0

    for m in module_metrics:
        module = str(m.get("logical_module") or m.get("module") or "unknown")
        item = per_module.setdefault(
            module,
            {
                "module": module,
                "attempts": 0,
                "success_calls": 0,
                "failed_calls": 0,
                "elapsed_seconds": 0.0,
                "prompt_tokens": 0,
                "completion_tokens": 0,
                "total_tokens": 0,
                "last_model": "",
                "last_status": "",
            },
        )
        item["attempts"] += 1
        item["last_model"] = str(m.get("model", ""))
        item["last_status"] = str(m.get("status", ""))
        elapsed = float(m.get("elapsed_seconds", 0.0) or 0.0)
        status = str(m.get("status", ""))
        if status == "success":
            p = int(m.get("prompt_tokens", 0) or 0)
            c = int(m.get("completion_tokens", 0) or 0)
            t = int(m.get("total_tokens", 0) or 0)
            item["success_calls"] += 1
            item["elapsed_seconds"] += elapsed
            item["prompt_tokens"] += p
            item["completion_tokens"] += c
            item["total_tokens"] += t
            total_prompt += p
            total_completion += c
            total_tokens += t
            total_elapsed += elapsed
            success_calls += 1
        else:
            item["failed_calls"] += 1
            fail_calls += 1

    modules = sorted(
        per_module.values(),
        key=lambda x: (float(x.get("elapsed_seconds", 0.0)), int(x.get("total_tokens", 0))),
        reverse=True,
    )
    return {
        "summary_generated_at": datetime.now().isoformat(),
        "api_calls_total": len(module_metrics),
        "api_calls_success": success_calls,
        "api_calls_failed": fail_calls,
        "total_prompt_tokens": total_prompt,
        "total_completion_tokens": total_completion,
        "total_tokens": total_tokens,
        "total_elapsed_seconds": round(total_elapsed, 3),
        "modules": modules,
    }


def module_metrics_markdown(summary: Dict[str, Any]) -> str:
    lines = [
        "# 模块耗时与Token统计",
        "",
        f"- API调用总数: {summary.get('api_calls_total', 0)}",
        f"- 成功调用: {summary.get('api_calls_success', 0)}",
        f"- 失败调用: {summary.get('api_calls_failed', 0)}",
        f"- 输入Token总计: {summary.get('total_prompt_tokens', 0)}",
        f"- 输出Token总计: {summary.get('total_completion_tokens', 0)}",
        f"- 总Token: {summary.get('total_tokens', 0)}",
        f"- 成功调用总耗时(秒): {summary.get('total_elapsed_seconds', 0)}",
        "",
        "| 模块 | 尝试次数 | 成功 | 失败 | 耗时(秒) | 输入Token | 输出Token | 总Token |",
        "|---|---:|---:|---:|---:|---:|---:|---:|",
    ]
    for m in summary.get("modules", []):
        lines.append(
            f"| {m.get('module','')} | {m.get('attempts',0)} | {m.get('success_calls',0)} | "
            f"{m.get('failed_calls',0)} | {float(m.get('elapsed_seconds',0.0)):.2f} | "
            f"{m.get('prompt_tokens',0)} | {m.get('completion_tokens',0)} | {m.get('total_tokens',0)} |"
        )
    return "\n".join(lines) + "\n"


def call_deepseek(
    client: OpenAI,
    model: str,
    system_prompt: str,
    user_prompt: str,
    module_name: str,
    module_logs: List[str],
    module_metrics: List[Dict[str, Any]] | None = None,
    temperature: float = 0.2,
    request_timeout: float = 60.0,
    max_retries: int = 3,
) -> str:
    last_error = ""
    for attempt in range(1, max_retries + 1):
        t0 = time.time()
        print(f"  - 模块 {module_name}: 尝试 {attempt}/{max_retries}", flush=True)
        module_logs.append(
            f"[{datetime.now().isoformat()}] {module_name}: start attempt={attempt}/{max_retries} timeout={request_timeout}s"
        )
        try:
            resp = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
                temperature=temperature,
                timeout=request_timeout,
            )
            content = (resp.choices[0].message.content or "").strip()
            elapsed = time.time() - t0
            if not content:
                raise RuntimeError("返回空内容")
            usage = extract_usage_info(resp)
            module_logs.append(
                f"[{datetime.now().isoformat()}] {module_name}: success attempt={attempt} "
                f"elapsed={elapsed:.2f}s len={len(content)} tokens={usage['total_tokens']}"
            )
            if module_metrics is not None:
                module_metrics.append(
                    {
                        "timestamp": datetime.now().isoformat(),
                        "module": module_name,
                        "logical_module": _logical_module_name(module_name),
                        "attempt": attempt,
                        "status": "success",
                        "model": model,
                        "elapsed_seconds": round(float(elapsed), 4),
                        "prompt_tokens": usage["prompt_tokens"],
                        "completion_tokens": usage["completion_tokens"],
                        "total_tokens": usage["total_tokens"],
                        "finish_reason": usage["finish_reason"],
                        "request_timeout": float(request_timeout),
                    }
                )
            print(f"    -> 成功 ({elapsed:.2f}s)", flush=True)
            return content
        except Exception as e:
            elapsed = time.time() - t0
            last_error = str(e)
            module_logs.append(
                f"[{datetime.now().isoformat()}] {module_name}: fail attempt={attempt} elapsed={elapsed:.2f}s err={last_error}"
            )
            if module_metrics is not None:
                module_metrics.append(
                    {
                        "timestamp": datetime.now().isoformat(),
                        "module": module_name,
                        "logical_module": _logical_module_name(module_name),
                        "attempt": attempt,
                        "status": "failed",
                        "model": model,
                        "elapsed_seconds": round(float(elapsed), 4),
                        "prompt_tokens": 0,
                        "completion_tokens": 0,
                        "total_tokens": 0,
                        "finish_reason": "",
                        "error": last_error[:500],
                        "request_timeout": float(request_timeout),
                    }
                )
            print(f"    -> 失败 ({elapsed:.2f}s): {last_error}", flush=True)
            if attempt < max_retries:
                time.sleep(min(2 * attempt, 5))

    raise RuntimeError(f"{module_name} 调用失败: {last_error}")


def generate_reports(
    client: OpenAI,
    model: str,
    final_model: str,
    mode: str,
    data_context: str,
    market_snapshot: Dict[str, Any],
    news: List[Dict[str, Any]],
    module_logs: List[str],
    module_metrics: List[Dict[str, Any]],
    request_timeout: float,
    max_retries: int,
    analyst_workers: int = 3,
    continue_on_error: bool = True,
    reused_reports: Dict[str, str] | None = None,
    reused_source_label: str = "",
) -> Tuple[Dict[str, str], Dict[str, Any]]:
    reports: Dict[str, str] = dict(reused_reports or {})
    degraded_modules: List[str] = []

    def run_module(
        module_name: str,
        system_prompt: str,
        user_prompt: str,
        temperature: float = 0.2,
        model_override: str | None = None,
        fallback_model: str | None = None,
        fallback_text: str | None = None,
    ) -> str:
        active_model = (model_override or model).strip()
        print(f"  - 生成 {module_name} (model={active_model})", flush=True)
        try:
            return call_deepseek(
                client=client,
                model=active_model,
                module_name=module_name,
                module_logs=module_logs,
                module_metrics=module_metrics,
                system_prompt=system_prompt,
                user_prompt=user_prompt,
                temperature=temperature,
                request_timeout=request_timeout,
                max_retries=max_retries,
            )
        except Exception as e:
            module_logs.append(
                f"[{datetime.now().isoformat()}] module_error name={module_name} model={active_model} err={e}"
            )
            if fallback_model and fallback_model != active_model:
                try:
                    module_logs.append(
                        f"[{datetime.now().isoformat()}] module_fallback_retry name={module_name} model={fallback_model}"
                    )
                    retry_text = call_deepseek(
                        client=client,
                        model=fallback_model,
                        module_name=f"{module_name}_fallback",
                        module_logs=module_logs,
                        module_metrics=module_metrics,
                        system_prompt=system_prompt,
                        user_prompt=user_prompt,
                        temperature=temperature,
                        request_timeout=request_timeout,
                        max_retries=max(1, min(2, max_retries)),
                    )
                    degraded_modules.append(module_name)
                    return (
                        f"> 注意：模块 `{module_name}` 主模型 `{active_model}` 失败，已降级到 `{fallback_model}`。\n\n"
                        + retry_text
                    )
                except Exception as e2:
                    module_logs.append(
                        f"[{datetime.now().isoformat()}] module_fallback_fail name={module_name} model={fallback_model} err={e2}"
                    )
                    if not continue_on_error:
                        raise
                    degraded_modules.append(module_name)
                    return fallback_text or (
                        f"# {module_name}\n\n"
                        "## 降级说明\n\n"
                        f"- 模块失败并降级输出\n- 错误：`{str(e2)[:300]}`\n\n"
                        "## 临时结论\n\n"
                        "该模块未获得完整结果，请参考其余模块并稍后重跑。\n"
                    )

            if not continue_on_error:
                raise
            degraded_modules.append(module_name)
            return fallback_text or (
                f"# {module_name}\n\n"
                "## 降级说明\n\n"
                f"- 模块失败并降级输出\n- 错误：`{str(e)[:300]}`\n\n"
                "## 临时结论\n\n"
                "该模块未获得完整结果，请参考其余模块并稍后重跑。\n"
            )

    # 1) 并行分析师模块（核心提速点）
    analyst_prompts = {
        "market_report": (
            "你是市场技术分析师。必须用中文Markdown输出，数据不足时明确写出。",
            (
                "请生成 market_report.md 内容，聚焦技术面和行情：\n"
                "- 均线、近期涨跌、成交量变化\n"
                "- 支撑位/压力位\n"
                "- 结尾给出技术面倾向（偏多/中性/偏空）\n\n"
                f"{data_context}"
            ),
        ),
        "fundamentals_report": (
            "你是基本面分析师。必须用中文Markdown输出。",
            (
                "请生成 fundamentals_report.md 内容：\n"
                "- 公司/标的性质（若是ETF请明确）\n"
                "- 估值与财务指标解读（PE/PB/市值/股息/beta等）\n"
                "- 指标缺失时请说明并给出保守判断\n"
                "- 给出中短期基本面结论\n\n"
                f"{data_context}"
            ),
        ),
        "news_report": (
            "你是新闻事件分析师。必须用中文Markdown输出。",
            (
                "请生成 news_report.md 内容：\n"
                "- 列出最近关键新闻（无新闻就写'暂无高相关新闻'）\n"
                "- 评估对价格可能影响（短期/中期）\n"
                "- 给出新闻面情绪判断（偏多/中性/偏空）\n\n"
                f"{data_context}"
            ),
        ),
    }

    workers = max(1, min(analyst_workers, 3))
    print(f"  - [并行] 分析师模块 workers={workers}", flush=True)
    with ThreadPoolExecutor(max_workers=workers) as pool:
        fut_map = {
            pool.submit(
                run_module,
                module_name=name,
                system_prompt=system_prompt,
                user_prompt=user_prompt,
            ): name
            for name, (system_prompt, user_prompt) in analyst_prompts.items()
            if f"{name}.md" not in reports
        }
        for fut in as_completed(fut_map):
            name = fut_map[fut]
            reports[f"{name}.md"] = fut.result()

    market_report = reports.get("market_report.md", "")
    fundamentals_report = reports.get("fundamentals_report.md", "")
    news_report = reports.get("news_report.md", "")

    if "research_team_decision.md" in reports:
        research_team_decision = reports["research_team_decision.md"]
        module_logs.append(
            f"[{datetime.now().isoformat()}] reuse_report name=research_team_decision source={reused_source_label}"
        )
    else:
        research_team_decision = run_module(
            module_name="research_team_decision",
            system_prompt="你是研究团队主持人。必须用中文Markdown输出。",
            user_prompt=(
                "请生成 research_team_decision.md，必须包含三部分：\n"
                "## 多头研究员观点\n"
                "## 空头研究员观点\n"
                "## 研究经理综合决策\n"
                "要求：论据来自行情/基本面/新闻，不要空话。\n\n"
                f"market_report:\n{market_report}\n\n"
                f"fundamentals_report:\n{fundamentals_report}\n\n"
                f"news_report:\n{news_report}\n"
            ),
        )
    reports["research_team_decision.md"] = research_team_decision

    if "investment_plan.md" in reports:
        investment_plan = reports["investment_plan.md"]
        module_logs.append(
            f"[{datetime.now().isoformat()}] reuse_report name=investment_plan source={reused_source_label}"
        )
    else:
        investment_plan = run_module(
            module_name="investment_plan",
            system_prompt="你是投资组合经理。必须用中文Markdown输出。",
            user_prompt=(
                "请生成 investment_plan.md：\n"
                "- 明确给出当前建议：买入/持有/卖出（三选一）\n"
                "- 给出仓位建议（如轻仓/中仓/空仓）\n"
                "- 给出触发条件（什么情况下改变观点）\n\n"
                f"research_team_decision:\n{research_team_decision}\n"
            ),
        )
    reports["investment_plan.md"] = investment_plan

    trader_investment_plan = run_module(
        module_name="trader_investment_plan",
        system_prompt="你是交易员。必须用中文Markdown输出。",
        user_prompt=(
            "请生成 trader_investment_plan.md：\n"
            "- 入场区间\n"
            "- 止损位\n"
            "- 第一/第二目标位\n"
            "- 执行节奏（分批、一次性等）\n"
            "- 如果是ETF请按指数型资产特点给计划\n\n"
            f"investment_plan:\n{investment_plan}\n\n"
            f"market_report:\n{market_report}\n"
        ),
    )
    reports["trader_investment_plan.md"] = trader_investment_plan

    if "risk_management_decision.md" in reports:
        risk_management_decision = reports["risk_management_decision.md"]
        module_logs.append(
            f"[{datetime.now().isoformat()}] reuse_report name=risk_management_decision source={reused_source_label}"
        )
    else:
        risk_management_decision = run_module(
            module_name="risk_management_decision",
            system_prompt="你是风险管理委员会。必须用中文Markdown输出。",
            user_prompt=(
                "请生成 risk_management_decision.md，必须包含：\n"
                "## 激进风险观点\n"
                "## 保守风险观点\n"
                "## 中性风险观点\n"
                "## 风险委员会最终裁决\n"
                "并明确主要风险来源和应对动作。\n\n"
                f"trader_investment_plan:\n{trader_investment_plan}\n\n"
                f"investment_plan:\n{investment_plan}\n\n"
                f"news_report:\n{news_report}\n"
            ),
        )
    reports["risk_management_decision.md"] = risk_management_decision

    # 结构化最终决策：可用 reasoner，失败自动降级到主模型
    raw_final_json = run_module(
        module_name="final_trade_decision_json",
        model_override=final_model,
        fallback_model=model if final_model != model else None,
        fallback_text=(
            '{\n'
            '  "action": "持有",\n'
            '  "confidence": 0.55,\n'
            '  "risk_score": 0.50,\n'
            '  "target_price_range": "待确认",\n'
            '  "reasoning": "最终决策模块失败，已按保守策略输出中性建议。"\n'
            '}'
        ),
        system_prompt="你是投资决策结构化助手，只返回JSON，不要返回markdown。",
        user_prompt=(
            "请仅返回一个JSON对象，字段必须完整：\n"
            "{\n"
            '  "action": "买入|持有|卖出",\n'
            '  "confidence": 0到1的小数,\n'
            '  "risk_score": 0到1的小数,\n'
            '  "target_price_range": "例如 4.50-5.20 CNY",\n'
            '  "reasoning": "不超过200字"\n'
            "}\n\n"
            f"market_report:\n{market_report}\n\n"
            f"fundamentals_report:\n{fundamentals_report}\n\n"
            f"news_report:\n{news_report}\n\n"
            f"investment_plan:\n{investment_plan}\n\n"
            f"risk_management_decision:\n{risk_management_decision}\n"
        ),
        temperature=0.0,
    )

    decision = safe_json_extract(raw_final_json) or {}

    action = str(decision.get("action", "持有"))
    if action not in {"买入", "持有", "卖出"}:
        action_lower = action.lower()
        if "buy" in action_lower:
            action = "买入"
        elif "sell" in action_lower:
            action = "卖出"
        else:
            action = "持有"

    def to_float(v: Any, default: float) -> float:
        try:
            x = float(v)
            if x < 0:
                return 0.0
            if x > 1:
                return 1.0
            return x
        except Exception:
            return default

    decision_obj = {
        "action": action,
        "confidence": to_float(decision.get("confidence", 0.68), 0.68),
        "risk_score": to_float(decision.get("risk_score", 0.42), 0.42),
        "target_price_range": str(decision.get("target_price_range", "待进一步确认")).strip(),
        "reasoning": str(decision.get("reasoning", "基于综合分析给出中性建议")).strip(),
        "mode": mode,
        "model_main": model,
        "model_final": final_model,
        "degraded_modules": degraded_modules,
        "reused_direction_source": reused_source_label,
    }

    final_trade_decision_md = (
        "# 最终交易决策\n\n"
        f"- **建议动作**: {decision_obj['action']}\n"
        f"- **置信度**: {decision_obj['confidence']:.2f}\n"
        f"- **风险评分**: {decision_obj['risk_score']:.2f}\n"
        f"- **目标区间**: {decision_obj['target_price_range']}\n"
        f"- **执行模式**: {decision_obj['mode']}（主模型: {decision_obj['model_main']} / 最终模型: {decision_obj['model_final']}）\n\n"
        "## 决策理由\n\n"
        f"{decision_obj['reasoning']}\n"
    )
    reports["final_trade_decision.md"] = final_trade_decision_md

    final_report = (
        "# 综合分析总报告\n\n"
        "## 模块目录\n"
        "1. 市场分析\n"
        "2. 基本面分析\n"
        "3. 新闻分析\n"
        "4. 研究团队决策\n"
        "5. 投资计划\n"
        "6. 交易执行计划\n"
        "7. 风险管理决策\n"
        "8. 最终交易决策\n\n"
        "---\n\n"
        "## 1. 市场分析\n\n"
        f"{market_report}\n\n"
        "---\n\n"
        "## 2. 基本面分析\n\n"
        f"{fundamentals_report}\n\n"
        "---\n\n"
        "## 3. 新闻分析\n\n"
        f"{news_report}\n\n"
        "---\n\n"
        "## 4. 研究团队决策\n\n"
        f"{research_team_decision}\n\n"
        "---\n\n"
        "## 5. 投资计划\n\n"
        f"{investment_plan}\n\n"
        "---\n\n"
        "## 6. 交易执行计划\n\n"
        f"{trader_investment_plan}\n\n"
        "---\n\n"
        "## 7. 风险管理决策\n\n"
        f"{risk_management_decision}\n\n"
        "---\n\n"
        "## 8. 最终交易决策\n\n"
        f"{final_trade_decision_md}\n"
    )
    if degraded_modules:
        final_report += (
            "\n## 降级记录\n\n"
            f"- 以下模块发生超时/错误并已自动降级继续：{'、'.join(degraded_modules)}\n"
            "- 建议网络稳定后再次运行以获得更完整结果。\n"
        )
    if reused_source_label:
        final_report += (
            "\n## 方向缓存复用\n\n"
            f"- 本次复用了较早分析的中期方向判断：{reused_source_label}\n"
            "- 当日重新刷新的部分主要是市场面、交易执行计划和最终交易决策。\n"
        )
    reports["final_report.md"] = final_report

    return reports, decision_obj


def normalize_report_markdown(text: str) -> str:
    """清理模型返回中常见的外层 markdown 代码块包裹，避免导出时显示异常。"""
    if not text:
        return ""
    cleaned = text.strip()
    # 去掉完整包裹的 ```markdown ... ``` 或 ``` ... ```
    fence_full = re.match(r"^```(?:markdown)?\s*([\s\S]*?)\s*```$", cleaned, flags=re.IGNORECASE)
    if fence_full:
        cleaned = fence_full.group(1).strip()
    # 去掉残留的单独 fence 行
    cleaned = re.sub(r"(?m)^\s*```(?:markdown)?\s*$", "", cleaned)
    cleaned = re.sub(r"(?m)^\s*```\s*$", "", cleaned)
    return cleaned.strip()


def build_share_markdown(
    user_symbol: str,
    analysis_date: str,
    decision_obj: Dict[str, Any],
    reports: Dict[str, str],
    market_snapshot: Dict[str, Any],
) -> str:
    latest = market_snapshot.get("latest", {})
    fundamentals = market_snapshot.get("fundamentals", {})
    technical = market_snapshot.get("technical", {})

    market_txt = normalize_report_markdown(reports.get("market_report.md", "无"))
    fundamentals_txt = normalize_report_markdown(reports.get("fundamentals_report.md", "无"))
    news_txt = normalize_report_markdown(reports.get("news_report.md", "无"))
    research_txt = normalize_report_markdown(reports.get("research_team_decision.md", "无"))
    invest_txt = normalize_report_markdown(reports.get("investment_plan.md", "无"))
    trader_txt = normalize_report_markdown(reports.get("trader_investment_plan.md", "无"))
    risk_txt = normalize_report_markdown(reports.get("risk_management_decision.md", "无"))
    final_txt = normalize_report_markdown(reports.get("final_trade_decision.md", "无"))

    return (
        f"# {user_symbol} 投资分析简报（可转发版）\n\n"
        f"> 分析日期：{analysis_date}\n\n"
        "## 一句话结论\n\n"
        f"- 建议动作：**{decision_obj.get('action', '持有')}**\n"
        f"- 置信度：**{decision_obj.get('confidence', 0):.2f}**\n"
        f"- 风险评分：**{decision_obj.get('risk_score', 0):.2f}**\n"
        f"- 目标区间：**{decision_obj.get('target_price_range', '待确认')}**\n\n"
        "## 核心理由\n\n"
        f"{decision_obj.get('reasoning', '无')}\n\n"
        "## 关键行情快照\n\n"
        f"- 最新收盘：{latest.get('close', 'N/A')}\n"
        f"- 最新成交量：{latest.get('volume', 'N/A')}\n"
        f"- 近1周涨跌：{technical.get('change_1w', 'N/A')}\n"
        f"- 近1月涨跌：{technical.get('change_1m', 'N/A')}\n"
        f"- 行业：{fundamentals.get('industry', 'N/A')}\n"
        f"- 市值：{fundamentals.get('marketCap', 'N/A')}\n\n"
        "## 详细分析（节选）\n\n"
        "### 市场面\n\n"
        f"{market_txt}\n\n"
        "### 基本面\n\n"
        f"{fundamentals_txt}\n\n"
        "### 新闻面\n\n"
        f"{news_txt}\n\n"
        "### 研究团队决策\n\n"
        f"{research_txt}\n\n"
        "### 投资计划\n\n"
        f"{invest_txt}\n\n"
        "### 交易执行计划\n\n"
        f"{trader_txt}\n\n"
        "### 风险管理\n\n"
        f"{risk_txt}\n\n"
        "### 最终交易决策\n\n"
        f"{final_txt}\n\n"
        "---\n\n"
        "风险提示：以上内容仅供学习交流，不构成投资建议。\n"
    )


def markdown_to_simple_html(md_text: str, title: str) -> str:
    try:
        import markdown as mdlib  # type: ignore

        body = mdlib.markdown(
            md_text,
            extensions=["extra", "tables", "fenced_code", "sane_lists", "nl2br"],
            output_format="html5",
        )
        return (
            "<!doctype html>\n"
            "<html lang='zh-CN'>\n"
            "<head>\n"
            "  <meta charset='utf-8'/>\n"
            "  <meta name='viewport' content='width=device-width, initial-scale=1'/>\n"
            f"  <title>{html.escape(title)}</title>\n"
            "  <style>\n"
            "    body{font-family:-apple-system,BlinkMacSystemFont,'PingFang SC','Microsoft YaHei',sans-serif;"
            "max-width:900px;margin:0 auto;padding:24px;line-height:1.75;color:#1f2937;background:#f7fafc;}\n"
            "    h1,h2,h3{line-height:1.35;color:#0f172a;}\n"
            "    h1{font-size:28px;margin-top:0;} h2{font-size:22px;margin-top:28px;} h3{font-size:18px;margin-top:20px;}\n"
            "    p,li{font-size:16px;} ul{padding-left:22px;} li{margin:6px 0;}\n"
            "    blockquote{background:#eef2ff;border-left:4px solid #6366f1;padding:10px 14px;border-radius:8px;}\n"
            "    code{background:#eef2f7;padding:2px 6px;border-radius:6px;}\n"
            "    pre{background:#0f172a;color:#e5e7eb;padding:14px;border-radius:10px;overflow:auto;}\n"
            "    table{border-collapse:collapse;width:100%;font-size:14px;margin:14px 0;}\n"
            "    th,td{border:1px solid #d1d5db;padding:8px;text-align:left;vertical-align:top;}\n"
            "    th{background:#f3f4f6;}\n"
            "    @media (max-width:768px){body{padding:14px;} h1{font-size:24px;} h2{font-size:20px;}}\n"
            "  </style>\n"
            "</head>\n"
            f"<body>{body}</body>\n"
            "</html>\n"
        )
    except Exception:
        pass

    lines = md_text.splitlines()
    html_parts: List[str] = []
    in_list = False

    def inline(text: str) -> str:
        rendered = html.escape(text)
        rendered = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", rendered)
        rendered = re.sub(r"`([^`]+)`", r"<code>\1</code>", rendered)
        return rendered

    for raw_line in lines:
        line = raw_line.rstrip()
        if not line:
            if in_list:
                html_parts.append("</ul>")
                in_list = False
            html_parts.append("<p></p>")
            continue

        if line == "---":
            if in_list:
                html_parts.append("</ul>")
                in_list = False
            html_parts.append("<hr/>")
            continue

        if line.startswith("# "):
            if in_list:
                html_parts.append("</ul>")
                in_list = False
            html_parts.append(f"<h1>{inline(line[2:])}</h1>")
            continue
        if line.startswith("## "):
            if in_list:
                html_parts.append("</ul>")
                in_list = False
            html_parts.append(f"<h2>{inline(line[3:])}</h2>")
            continue
        if line.startswith("### "):
            if in_list:
                html_parts.append("</ul>")
                in_list = False
            html_parts.append(f"<h3>{inline(line[4:])}</h3>")
            continue

        if line.startswith("> "):
            if in_list:
                html_parts.append("</ul>")
                in_list = False
            html_parts.append(f"<blockquote>{inline(line[2:])}</blockquote>")
            continue

        if line.startswith("- "):
            if not in_list:
                html_parts.append("<ul>")
                in_list = True
            html_parts.append(f"<li>{inline(line[2:])}</li>")
            continue

        if in_list:
            html_parts.append("</ul>")
            in_list = False
        html_parts.append(f"<p>{inline(line)}</p>")

    if in_list:
        html_parts.append("</ul>")

    body = "\n".join(html_parts)
    return (
        "<!doctype html>\n"
        "<html lang='zh-CN'>\n"
        "<head>\n"
        "  <meta charset='utf-8'/>\n"
        "  <meta name='viewport' content='width=device-width, initial-scale=1'/>\n"
        f"  <title>{html.escape(title)}</title>\n"
        "  <style>\n"
        "    body{font-family:-apple-system,BlinkMacSystemFont,'PingFang SC','Microsoft YaHei',sans-serif;"
        "max-width:900px;margin:0 auto;padding:24px;line-height:1.75;color:#1f2937;background:#f7fafc;}\n"
        "    h1,h2,h3{line-height:1.35;color:#0f172a;}\n"
        "    h1{font-size:28px;margin-top:0;} h2{font-size:22px;margin-top:28px;} h3{font-size:18px;margin-top:20px;}\n"
        "    p,li{font-size:16px;} ul{padding-left:22px;} li{margin:6px 0;}\n"
        "    blockquote{background:#eef2ff;border-left:4px solid #6366f1;padding:10px 14px;border-radius:8px;}\n"
        "    @media (max-width:768px){body{padding:14px;} h1{font-size:24px;} h2{font-size:20px;}}\n"
        "  </style>\n"
        "</head>\n"
        f"<body>{body}</body>\n"
        "</html>\n"
    )


def markdown_to_html_fragment(md_text: str) -> str:
    cleaned = normalize_report_markdown(md_text)
    if not cleaned:
        return "<p>无内容</p>"
    try:
        import markdown as mdlib  # type: ignore

        return mdlib.markdown(
            cleaned,
            extensions=["extra", "tables", "fenced_code", "sane_lists", "nl2br"],
            output_format="html5",
        )
    except Exception:
        return "<p>" + html.escape(cleaned).replace("\n", "<br/>") + "</p>"


def build_share_html(
    user_symbol: str,
    analysis_date: str,
    decision_obj: Dict[str, Any],
    reports: Dict[str, str],
    market_snapshot: Dict[str, Any],
) -> str:
    latest = market_snapshot.get("latest", {})
    technical = market_snapshot.get("technical", {})
    fundamentals = market_snapshot.get("fundamentals", {})

    action = str(decision_obj.get("action", "持有"))
    action_class = {
        "买入": "buy",
        "持有": "hold",
        "卖出": "sell",
    }.get(action, "hold")

    def fmt_pct(v: Any) -> str:
        try:
            return f"{float(v) * 100:.2f}%"
        except Exception:
            return "N/A"

    modules = [
        ("市场分析", reports.get("market_report.md", "")),
        ("基本面分析", reports.get("fundamentals_report.md", "")),
        ("新闻分析", reports.get("news_report.md", "")),
        ("研究团队决策", reports.get("research_team_decision.md", "")),
        ("投资计划", reports.get("investment_plan.md", "")),
        ("交易执行计划", reports.get("trader_investment_plan.md", "")),
        ("风险管理", reports.get("risk_management_decision.md", "")),
        ("最终交易决策", reports.get("final_trade_decision.md", "")),
    ]

    module_blocks = []
    for title, text in modules:
        module_blocks.append(
            "<details class='module-card'>"
            f"<summary>{html.escape(title)}</summary>"
            f"<div class='module-body'>{markdown_to_html_fragment(text)}</div>"
            "</details>"
        )

    return (
        "<!doctype html><html lang='zh-CN'><head>"
        "<meta charset='utf-8'/>"
        "<meta name='viewport' content='width=device-width, initial-scale=1'/>"
        f"<title>{html.escape(user_symbol)} 投资分析简报</title>"
        "<style>"
        ":root{--bg:#f4f7fb;--card:#fff;--text:#0f172a;--muted:#64748b;--border:#dbe3ee;"
        "--buy:#16a34a;--hold:#f59e0b;--sell:#dc2626;--primary:#2563eb;}"
        "body{margin:0;background:linear-gradient(165deg,#eef3ff,#f8fbff 38%,#f4f7fb);color:var(--text);"
        "font-family:-apple-system,BlinkMacSystemFont,'PingFang SC','Microsoft YaHei',sans-serif;}"
        ".wrap{max-width:980px;margin:0 auto;padding:18px;}"
        ".hero{background:var(--card);border:1px solid var(--border);border-radius:16px;padding:16px 18px;"
        "box-shadow:0 12px 24px rgba(15,23,42,.08);}"
        ".top{display:flex;justify-content:space-between;gap:12px;align-items:center;flex-wrap:wrap;}"
        ".title{font-size:28px;font-weight:800;margin:0;}"
        ".sub{margin:6px 0 0;color:var(--muted);font-size:14px;}"
        ".badge{padding:7px 14px;border-radius:999px;color:#fff;font-weight:700;font-size:14px;}"
        ".badge.buy{background:var(--buy);} .badge.hold{background:var(--hold);} .badge.sell{background:var(--sell);}"
        ".kpi{display:grid;grid-template-columns:repeat(4,minmax(120px,1fr));gap:10px;margin-top:12px;}"
        ".kpi .item{background:#f8fbff;border:1px solid var(--border);border-radius:12px;padding:10px;}"
        ".kpi .label{font-size:12px;color:var(--muted);} .kpi .val{font-size:18px;font-weight:700;margin-top:4px;}"
        ".reason{margin-top:12px;background:#f8fbff;border:1px solid var(--border);border-radius:12px;padding:12px;line-height:1.75;}"
        ".module-list{margin-top:14px;display:grid;gap:10px;}"
        ".module-card{background:var(--card);border:1px solid var(--border);border-radius:12px;padding:0 12px;}"
        ".module-card > summary{cursor:pointer;list-style:none;padding:13px 2px;font-size:17px;font-weight:700;}"
        ".module-card > summary::-webkit-details-marker{display:none;}"
        ".module-body{border-top:1px dashed var(--border);padding:12px 2px 14px;line-height:1.75;}"
        ".module-body table{border-collapse:collapse;width:100%;font-size:14px;margin:10px 0;}"
        ".module-body th,.module-body td{border:1px solid #d4dbe5;padding:8px;text-align:left;vertical-align:top;}"
        ".module-body th{background:#f3f6fb;}"
        ".module-body code{background:#eef2f7;padding:2px 5px;border-radius:6px;}"
        ".module-body pre{overflow:auto;background:#0f172a;color:#e2e8f0;padding:12px;border-radius:10px;}"
        ".foot{margin:16px 0 4px;color:var(--muted);font-size:13px;text-align:center;}"
        "@media (max-width:840px){.title{font-size:23px;}.kpi{grid-template-columns:repeat(2,minmax(120px,1fr));}}"
        "</style></head><body><div class='wrap'>"
        "<section class='hero'>"
        "<div class='top'>"
        f"<div><h1 class='title'>{html.escape(user_symbol)} 投资分析简报</h1>"
        f"<p class='sub'>分析日期：{html.escape(analysis_date)} | 模式：{html.escape(str(decision_obj.get('mode', 'quick')))}</p></div>"
        f"<span class='badge {action_class}'>建议：{html.escape(action)}</span>"
        "</div>"
        "<div class='kpi'>"
        f"<div class='item'><div class='label'>置信度</div><div class='val'>{float(decision_obj.get('confidence', 0)):.2f}</div></div>"
        f"<div class='item'><div class='label'>风险评分</div><div class='val'>{float(decision_obj.get('risk_score', 0)):.2f}</div></div>"
        f"<div class='item'><div class='label'>近1周涨跌</div><div class='val'>{fmt_pct(technical.get('change_1w'))}</div></div>"
        f"<div class='item'><div class='label'>近1月涨跌</div><div class='val'>{fmt_pct(technical.get('change_1m'))}</div></div>"
        "</div>"
        "<div class='kpi'>"
        f"<div class='item'><div class='label'>最新收盘</div><div class='val'>{html.escape(str(latest.get('close', 'N/A')))}</div></div>"
        f"<div class='item'><div class='label'>目标区间</div><div class='val'>{html.escape(str(decision_obj.get('target_price_range', '待确认')))}</div></div>"
        f"<div class='item'><div class='label'>行业</div><div class='val'>{html.escape(str(fundamentals.get('industry', 'N/A')))}</div></div>"
        f"<div class='item'><div class='label'>市值</div><div class='val'>{html.escape(str(fundamentals.get('marketCap', 'N/A')))}</div></div>"
        "</div>"
        f"<div class='reason'><strong>核心理由：</strong>{html.escape(str(decision_obj.get('reasoning', '无')))}</div>"
        "</section>"
        "<section class='module-list'>"
        + "".join(module_blocks) +
        "</section>"
        "<p class='foot'>风险提示：仅供学习交流，不构成投资建议。</p>"
        "</div></body></html>"
    )


def export_docx_from_markdown(md_text: str, out_path: Path) -> bool:
    try:
        from docx import Document  # type: ignore
    except Exception:
        return False

    doc = Document()
    for raw_line in md_text.splitlines():
        line = raw_line.strip()
        if not line:
            doc.add_paragraph("")
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            continue
        if line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
            continue
        doc.add_paragraph(line)

    doc.save(str(out_path))
    return True


def write_outputs(
    results_dir: Path,
    user_symbol: str,
    analysis_date: str,
    yf_symbol: str,
    model: str,
    market_type: str,
    reports: Dict[str, str],
    market_snapshot: Dict[str, Any],
    news: List[Dict[str, Any]],
    decision_obj: Dict[str, Any],
    module_metrics: List[Dict[str, Any]],
    module_logs: List[str],
) -> Path:
    stock_root = results_dir / user_symbol / analysis_date
    reports_dir = stock_root / "reports"
    reports_dir.mkdir(parents=True, exist_ok=True)

    for filename, content in reports.items():
        (reports_dir / filename).write_text(content, encoding="utf-8")

    (reports_dir / "market_snapshot.json").write_text(
        json.dumps(market_snapshot, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    (reports_dir / "news_snapshot.json").write_text(
        json.dumps(news, ensure_ascii=False, indent=2), encoding="utf-8"
    )

    (stock_root / "decision.json").write_text(
        json.dumps(decision_obj, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    metrics_summary = summarize_module_metrics(module_metrics)
    (stock_root / "module_metrics.json").write_text(
        json.dumps(module_metrics, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    (stock_root / "module_metrics_summary.json").write_text(
        json.dumps(metrics_summary, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    (reports_dir / "module_metrics.md").write_text(
        module_metrics_markdown(metrics_summary), encoding="utf-8"
    )

    report_types = [
        "market_report",
        "fundamentals_report",
        "news_report",
        "research_team_decision",
        "investment_plan",
        "trader_investment_plan",
        "risk_management_decision",
        "final_trade_decision",
        "module_metrics",
    ]

    metadata = {
        "stock_symbol": user_symbol,
        "normalized_symbol": yf_symbol,
        "analysis_date": analysis_date,
        "timestamp": datetime.now().isoformat(),
        "status": "completed",
        "market_type": market_type,
        "research_depth": 3 if decision_obj.get("mode") == "deep" else 1,
        "analysts": ["market", "fundamentals", "news"],
        "model": model,
        "mode": decision_obj.get("mode", "quick"),
        "model_main": decision_obj.get("model_main", model),
        "model_final": decision_obj.get("model_final", model),
        "degraded_modules": decision_obj.get("degraded_modules", []),
        "analysis_strategy": decision_obj.get("analysis_strategy", "full_analysis"),
        "reused_direction_source": decision_obj.get("reused_direction_source", ""),
        "generator": "minimal_deepseek_report.py",
        "reports_count": len(report_types),
        "report_types": report_types,
        "metrics_files": ["module_metrics.json", "module_metrics_summary.json"],
    }

    (stock_root / "analysis_metadata.json").write_text(
        json.dumps(metadata, ensure_ascii=False, indent=2), encoding="utf-8"
    )

    # 便于微信/IM转发的简版输出
    share_dir = stock_root / "share"
    share_dir.mkdir(parents=True, exist_ok=True)
    share_md = build_share_markdown(
        user_symbol=user_symbol,
        analysis_date=analysis_date,
        decision_obj=decision_obj,
        reports=reports,
        market_snapshot=market_snapshot,
    )
    share_base = f"{user_symbol}_{analysis_date}_share"
    share_md_path = share_dir / f"{share_base}.md"
    share_html_path = share_dir / f"{share_base}.html"
    share_txt_path = share_dir / f"{share_base}.txt"
    share_docx_path = share_dir / f"{share_base}.docx"

    share_md_path.write_text(share_md, encoding="utf-8")
    share_txt_path.write_text(share_md, encoding="utf-8")
    share_html_path.write_text(
        build_share_html(
            user_symbol=user_symbol,
            analysis_date=analysis_date,
            decision_obj=decision_obj,
            reports=reports,
            market_snapshot=market_snapshot,
        ),
        encoding="utf-8",
    )

    if export_docx_from_markdown(share_md, share_docx_path):
        module_logs.append(f"[{datetime.now().isoformat()}] share_docx ok path={share_docx_path}")
    else:
        module_logs.append(f"[{datetime.now().isoformat()}] share_docx skip reason=python-docx not installed")

    # 兼容原项目目录习惯（在最后写，保证包含完整日志）
    (stock_root / "message_tool.log").write_text("\n".join(module_logs) + "\n", encoding="utf-8")

    return stock_root


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="DeepSeek 单模型多模块股票分析工具")
    parser.add_argument("stock_symbol", help="股票代码，如 600028 / 000630 / AAPL / 0700")
    parser.add_argument(
        "--date",
        dest="analysis_date",
        default=datetime.now().strftime("%Y-%m-%d"),
        help="分析日期，默认今天 YYYY-MM-DD",
    )
    parser.add_argument(
        "--model",
        default=os.getenv("DEEPSEEK_MODEL", "deepseek-chat"),
        help="DeepSeek模型，可选 deepseek-chat / deepseek-reasoner，默认 deepseek-chat",
    )
    parser.add_argument(
        "--mode",
        default=os.getenv("TA_MIN_MODE", "quick"),
        choices=["quick", "deep"],
        help="分析模式：quick（默认，快）/ deep（深度）",
    )
    parser.add_argument(
        "--final-model",
        default=os.getenv("DEEPSEEK_FINAL_MODEL", ""),
        help="最终决策模型，不传则 quick=跟随主模型，deep=deepseek-reasoner",
    )
    parser.add_argument("--api-key", default="", help="DeepSeek API Key，不传则读取 DEEPSEEK_API_KEY")
    parser.add_argument(
        "--base-url",
        default=os.getenv("DEEPSEEK_BASE_URL", "https://api.deepseek.com"),
        help="DeepSeek Base URL",
    )
    parser.add_argument(
        "--results-dir",
        default=os.getenv("TRADINGAGENTS_RESULTS_DIR", "results"),
        help="输出根目录",
    )
    parser.add_argument("--max-news", type=int, default=8, help="新闻条数上限")
    parser.add_argument(
        "--request-timeout",
        type=float,
        default=None,
        help="单次模型请求超时秒数；未设置时 quick=45, deep=120",
    )
    parser.add_argument(
        "--retries",
        type=int,
        default=None,
        help="单模块重试次数；未设置时 quick=1, deep=2",
    )
    parser.add_argument(
        "--analyst-workers",
        type=int,
        default=int(os.getenv("TA_MIN_ANALYST_WORKERS", "3")),
        help="分析师并发数（1-3），默认3",
    )
    parser.add_argument(
        "--cache-ttl-hours",
        type=float,
        default=float(os.getenv("TA_MIN_CACHE_TTL_HOURS", "12")),
        help="行情/新闻缓存有效期（小时），默认12",
    )
    parser.add_argument(
        "--disable-cache",
        action="store_true",
        help="禁用行情/新闻缓存",
    )
    parser.add_argument(
        "--direction-cache-days",
        type=int,
        default=int(os.getenv("TA_MIN_DIRECTION_CACHE_DAYS", "3")),
        help="方向缓存复用天数，默认3天",
    )
    parser.add_argument(
        "--force-full-analysis",
        action="store_true",
        help="强制全量重跑，不复用同日结果和方向缓存",
    )
    parser.add_argument(
        "--strict",
        action="store_true",
        help="严格模式：任一模块失败即退出（默认关闭，默认会自动降级继续）",
    )
    return parser.parse_args()


def main() -> int:
    args = parse_args()

    model_alias = {
        "chat": "deepseek-chat",
        "deepseek_chat": "deepseek-chat",
        "reasoner": "deepseek-reasoner",
        "deepseek_reasoner": "deepseek-reasoner",
    }
    args.model = model_alias.get(args.model.strip().lower(), args.model.strip())
    args.mode = (args.mode or "quick").strip().lower()
    if args.mode not in {"quick", "deep"}:
        args.mode = "quick"
    args.final_model = model_alias.get(args.final_model.strip().lower(), args.final_model.strip()) if args.final_model else ""
    if not args.final_model:
        args.final_model = "deepseek-reasoner" if args.mode == "deep" else args.model

    if args.request_timeout is None:
        args.request_timeout = 120.0 if args.mode == "deep" else 45.0
    if args.retries is None:
        args.retries = 2 if args.mode == "deep" else 1

    try:
        datetime.strptime(args.analysis_date, "%Y-%m-%d")
    except ValueError:
        raise SystemExit("--date 格式错误，应为 YYYY-MM-DD")

    project_root = Path(__file__).resolve().parents[1]
    results_dir = Path(args.results_dir)
    if not results_dir.is_absolute():
        results_dir = project_root / results_dir
    cache_dir = results_dir / "_cache"
    cache_ttl_seconds = max(0, int(args.cache_ttl_hours * 3600))
    ensure_dir(cache_dir)

    user_symbol, yf_symbol = normalize_symbol(args.stock_symbol)

    module_logs: List[str] = []
    module_metrics: List[Dict[str, Any]] = []
    module_logs.append(
        f"[{datetime.now().isoformat()}] start symbol={user_symbol} preferred={yf_symbol} "
        f"date={args.analysis_date} mode={args.mode} model={args.model} final_model={args.final_model} "
        f"timeout={args.request_timeout} retries={args.retries} cache={'off' if args.disable_cache else 'on'} ttl_s={cache_ttl_seconds} "
        f"direction_cache_days={args.direction_cache_days} force_full={args.force_full_analysis}"
    )

    reused_result = None
    reused_reports: Dict[str, str] = {}
    reused_source_label = ""
    if not args.force_full_analysis:
        reused_result = find_recent_reusable_result(
            results_dir=results_dir,
            user_symbol=user_symbol,
            analysis_date=args.analysis_date,
            direction_cache_days=max(0, int(args.direction_cache_days)),
        )
        if reused_result and reused_result["type"] == "same_day":
            decision_path = reused_result["result_dir"] / "decision.json"
            decision_obj = _load_json_file(decision_path)
            print("[cache] 命中同日有效结果，直接复用，无需重跑")
            print(f"目录: {reused_result['result_dir']}")
            print(f"最终建议: {decision_obj.get('action')} | 目标区间: {decision_obj.get('target_price_range')}")
            return 0
        if reused_result and reused_result["type"] == "direction_cache":
            reused_reports = load_reused_reports(reused_result["result_dir"])
            reused_source_label = "{0}（距今{1}天）".format(
                reused_result["analysis_date"],
                reused_result["delta_days"],
            )
            print(
                "[cache] 命中方向缓存，复用慢变量分析: {0}".format(reused_source_label)
            )
            module_logs.append(
                f"[{datetime.now().isoformat()}] direction_cache hit source={reused_result['result_dir']}"
            )

    api_key = args.api_key.strip() or os.getenv("DEEPSEEK_API_KEY", "").strip()
    if not api_key:
        raise SystemExit("未检测到 DeepSeek API Key。请设置 DEEPSEEK_API_KEY 或传 --api-key")

    # 1) 行情
    last_error: Exception | None = None
    symbol_used = yf_symbol
    symbols_to_try = candidate_yf_symbols(user_symbol, yf_symbol)

    market_snapshot: Dict[str, Any] | None = None
    per_symbol_retries = 2
    for idx, sym in enumerate(symbols_to_try, start=1):
        for retry in range(1, per_symbol_retries + 1):
            print(
                f"[1/6] 获取行情数据: {sym} (候选 {idx}/{len(symbols_to_try)}, 重试 {retry}/{per_symbol_retries})"
            )
            try:
                market_snapshot = fetch_market_data_cached(
                    yf_symbol=sym,
                    as_of_date=args.analysis_date,
                    cache_dir=cache_dir,
                    results_dir=results_dir,
                    user_symbol=user_symbol,
                    module_logs=module_logs,
                    cache_ttl_seconds=cache_ttl_seconds,
                    disable_cache=args.disable_cache,
                )
                symbol_used = sym
                break
            except Exception as e:
                last_error = e
                module_logs.append(
                    f"[{datetime.now().isoformat()}] market_fetch fail symbol={sym} retry={retry}/{per_symbol_retries} err={e}"
                )
                print(f"  - 失败: {e}")
                if retry < per_symbol_retries:
                    time.sleep(2)
        if market_snapshot is not None:
            break

    if market_snapshot is None:
        failure_hint = (
            "这通常不是股票本身异常，而是东财实时接口临时不可用，且本地缓存/历史结果也不足以兜底。"
            if any(is_a_share_symbol(sym) for sym in symbols_to_try)
            else "这通常不是股票本身异常，而是 yfinance/Yahoo 对当前代码覆盖不稳定或被限流。"
        )
        raise SystemExit(
            "所有候选代码都获取失败: {0}; 最后错误: {1}; {2}"
            .format(symbols_to_try, last_error, failure_hint)
        )

    market_type = infer_market_type(symbol_used)

    # 2) 新闻
    print(f"[2/6] 获取新闻数据: {symbol_used}")
    news = fetch_news_cached(
        yf_symbol=symbol_used,
        max_news=args.max_news,
        cache_dir=cache_dir,
        module_logs=module_logs,
        cache_ttl_seconds=cache_ttl_seconds,
        disable_cache=args.disable_cache,
    )
    module_logs.append(f"[{datetime.now().isoformat()}] news_fetch ok count={len(news)}")

    # 3) DeepSeek多模块生成
    print(
        f"[3/6] 调用 DeepSeek 生成多模块报告 "
        f"(mode={args.mode}, model={args.model}, final_model={args.final_model}, "
        f"timeout={args.request_timeout}s, retries={args.retries})"
    )
    client = OpenAI(api_key=api_key, base_url=args.base_url, max_retries=0)

    data_context = build_data_context(
        user_symbol=user_symbol,
        yf_symbol=symbol_used,
        analysis_date=args.analysis_date,
        market_type=market_type,
        market_snapshot=market_snapshot,
        news=news,
    )

    reports, decision_obj = generate_reports(
        client=client,
        model=args.model,
        final_model=args.final_model,
        mode=args.mode,
        data_context=data_context,
        market_snapshot=market_snapshot,
        news=news,
        module_logs=module_logs,
        module_metrics=module_metrics,
        request_timeout=args.request_timeout,
        max_retries=args.retries,
        analyst_workers=args.analyst_workers,
        continue_on_error=not args.strict,
        reused_reports=reused_reports,
        reused_source_label=reused_source_label,
    )
    metrics_summary = summarize_module_metrics(module_metrics)
    decision_obj["module_metrics_summary"] = metrics_summary
    decision_obj["analysis_strategy"] = "direction_cache_refresh" if reused_source_label else "full_analysis"
    decision_obj["direction_cache_days"] = int(args.direction_cache_days)
    decision_obj["force_full_analysis"] = bool(args.force_full_analysis)

    # 4) 写文件
    print("[4/6] 写入结果目录")
    out_dir = write_outputs(
        results_dir=results_dir,
        user_symbol=user_symbol,
        analysis_date=args.analysis_date,
        yf_symbol=symbol_used,
        model=args.model,
        market_type=market_type,
        reports=reports,
        market_snapshot=market_snapshot,
        news=news,
        decision_obj=decision_obj,
        module_metrics=module_metrics,
        module_logs=module_logs,
    )

    # 5) 输出摘要
    print("[5/6] 生成摘要")
    print(f"  - 最终建议: {decision_obj.get('action')} | 目标区间: {decision_obj.get('target_price_range')}")
    print(
        "  - Token统计: "
        f"in={metrics_summary.get('total_prompt_tokens', 0)} "
        f"out={metrics_summary.get('total_completion_tokens', 0)} "
        f"total={metrics_summary.get('total_tokens', 0)} "
        f"calls={metrics_summary.get('api_calls_total', 0)}"
    )
    top_modules = metrics_summary.get("modules", [])[:5]
    if top_modules:
        top_str = ", ".join(
            f"{m.get('module')}={float(m.get('elapsed_seconds', 0.0)):.1f}s/{int(m.get('total_tokens', 0))}tok"
            for m in top_modules
        )
        print(f"  - 模块耗时TOP: {top_str}")
    if decision_obj.get("degraded_modules"):
        print(f"  - 自动降级模块: {', '.join(decision_obj.get('degraded_modules', []))}")

    # 6) 完成
    print("[6/6] 完成")
    print("\n✅ 生成完成")
    print(f"目录: {out_dir}")
    print(f"主报告: {out_dir / 'reports' / 'final_report.md'}")
    print(f"最终决策: {out_dir / 'reports' / 'final_trade_decision.md'}")
    share_base = f"{user_symbol}_{args.analysis_date}_share"
    print(f"转发简版(MD): {out_dir / 'share' / f'{share_base}.md'}")
    print(f"转发简版(HTML): {out_dir / 'share' / f'{share_base}.html'}")
    print(f"转发简版(Word): {out_dir / 'share' / f'{share_base}.docx'}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
